"""
Resume-Job Matcher Web Application
Flask-based interface for ML resume matching with AI optimization
"""

import os
import pickle
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import torch
from transformers import BertTokenizer, BertModel
from sentence_transformers import SentenceTransformer, util
import anthropic
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import PyPDF2
import nltk
from io import BytesIO
import json

# Import configuration
import config

# CRITICAL: Download NLTK data BEFORE importing inference_utils
# This prevents the "Resource stopwords not found" error
print("📥 Checking NLTK data...")
try:
    nltk.data.find('corpora/stopwords')
    nltk.data.find('corpora/wordnet')
    print("✅ NLTK data already present")
except LookupError:
    print("📥 Downloading NLTK data (first time only)...")
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
    nltk.download('omw-1.4', quiet=True)
    print("✅ NLTK data downloaded")

# Import inference utilities (AFTER downloading NLTK data)
import inference_utils

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def convert_numpy_types(obj):
    """
    Recursively convert NumPy types to Python native types for JSON serialization
    """
    import numpy as np
    
    if isinstance(obj, np.bool_):
        return bool(obj)
    elif isinstance(obj, np.integer):
        return int(obj)
    elif isinstance(obj, np.floating):
        return float(obj)
    elif isinstance(obj, np.ndarray):
        return obj.tolist()
    elif isinstance(obj, dict):
        return {key: convert_numpy_types(value) for key, value in obj.items()}
    elif isinstance(obj, list):
        return [convert_numpy_types(item) for item in obj]
    else:
        return obj

# ============================================================================
# FLASK APP INITIALIZATION
# ============================================================================
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = config.UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = config.MAX_FILE_SIZE

# Create upload folder if it doesn't exist
os.makedirs(config.UPLOAD_FOLDER, exist_ok=True)

# ============================================================================
# GLOBAL VARIABLES FOR MODELS
# ============================================================================
# These will be loaded once at startup
tfidf_vectorizer_combined = None
svm_combined = None
label_encoder = None
bert_model = None
tokenizer = None
ml_knowledge_classifier = None
skill_similarity_model = None

# ============================================================================
# MODEL LOADING
# ============================================================================

def load_models():
    """Load all ML models at startup"""
    global tfidf_vectorizer_combined, svm_combined, label_encoder
    global bert_model, tokenizer, ml_knowledge_classifier, skill_similarity_model
    
    print("\n" + "="*70)
    print("LOADING ML MODELS")
    print("="*70)
    
    # Load pickle models
    print("\n📦 Loading pickle models...")
    try:
        with open(config.TFIDF_VECTORIZER_PATH, 'rb') as f:
            tfidf_vectorizer_combined = pickle.load(f)
        print("  ✅ TF-IDF Vectorizer loaded")
        
        with open(config.SVM_MODEL_PATH, 'rb') as f:
            svm_combined = pickle.load(f)
        print("  ✅ SVM Model loaded")
        
        with open(config.LABEL_ENCODER_PATH, 'rb') as f:
            label_encoder = pickle.load(f)
        print("  ✅ Label Encoder loaded")
    except FileNotFoundError as e:
        print(f"\n❌ ERROR: Model file not found!")
        print(f"   {e}")
        print("\n📁 Please ensure these files are in the 'models/' folder:")
        print("   - tfidf_vectorizer_combined.pkl")
        print("   - svm_combined.pkl")
        print("   - label_encoder.pkl")
        raise
    
    # Load HuggingFace models
    print("\n🤗 Loading HuggingFace models...")
    print("  (This may take a few minutes on first run...)")
    
    try:
        # BERT
        print("\n  Loading BERT...")
        tokenizer = BertTokenizer.from_pretrained('bert-base-uncased')
        bert_model = BertModel.from_pretrained('bert-base-uncased')
        bert_model.eval()
        print("  ✅ BERT loaded")
        
        # JobBERT for skill extraction
        print("\n  Loading JobBERT...")
        from transformers import pipeline
        ml_knowledge_classifier = pipeline(
            "text-classification",
            model="jjzha/jobbert_knowledge_extraction",
            return_all_scores=False
        )
        print("  ✅ JobBERT loaded")
        
        # Sentence Transformer for skill similarity
        print("\n  Loading Sentence Transformer...")
        skill_similarity_model = SentenceTransformer('sentence-transformers/all-MiniLM-L6-v2')
        print("  ✅ Sentence Transformer loaded")
        
    except Exception as e:
        print(f"\n❌ ERROR loading HuggingFace models: {e}")
        raise
    
    print("\n" + "="*70)
    print("✅ ALL MODELS LOADED SUCCESSFULLY")
    print("="*70 + "\n")
    
    # Update inference_utils with loaded models
    inference_utils.tfidf_vectorizer_combined = tfidf_vectorizer_combined
    inference_utils.svm_combined = svm_combined
    inference_utils.label_encoder = label_encoder
    inference_utils.bert_model = bert_model
    inference_utils.tokenizer = tokenizer
    inference_utils.ml_knowledge_classifier = ml_knowledge_classifier
    inference_utils.skill_similarity_model = skill_similarity_model
    inference_utils.util = util

# ============================================================================
# FILE HANDLING UTILITIES
# ============================================================================

def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in config.ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    """Extract text from PDF file"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")

def extract_text_from_txt(file_path):
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        raise Exception(f"Error reading TXT: {str(e)}")

def process_resume_input(resume_text, resume_file):
    """Process resume from either text input or file upload"""
    if resume_file and resume_file.filename:
        # File upload
        if not allowed_file(resume_file.filename):
            raise ValueError("Invalid file type. Allowed: .txt, .pdf, .docx")
        
        filename = secure_filename(resume_file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        resume_file.save(filepath)
        
        # Extract text based on file type
        ext = filename.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            text = extract_text_from_pdf(filepath)
        elif ext == 'docx':
            text = extract_text_from_docx(filepath)
        else:  # txt
            text = extract_text_from_txt(filepath)
        
        # Clean up uploaded file
        os.remove(filepath)
        return text
    
    elif resume_text and resume_text.strip():
        # Text input
        return resume_text.strip()
    
    else:
        raise ValueError("Please provide either resume text or upload a file")

# ============================================================================
# CLAUDE AI INTEGRATION
# ============================================================================

def get_ai_suggestions(resume_text, job_description, missing_skills, matched_skills, score_details):
    """Get AI-powered suggestions for resume improvement using Claude"""
    try:
        client = anthropic.Anthropic(api_key=config.CLAUDE_API_KEY)
        
        # Create comprehensive prompt
        prompt = f"""You are an expert career advisor and resume optimization specialist. Analyze this resume and provide specific, actionable improvement suggestions.

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}

CURRENT MATCH ANALYSIS:
- Overall Score: {score_details['final_score']:.1%}
- BERT Semantic Similarity: {score_details['bert_similarity']:.1%}
- Keywords Match: {score_details['tfidf_similarity']:.1%}
- Category: Resume ({score_details['resume_category']}) vs Job ({score_details['job_category']})

MATCHED SKILLS ({len(matched_skills)}):
{', '.join([skill['resume_has'] for skill in matched_skills[:10]]) if matched_skills else 'None'}

MISSING SKILLS ({len(missing_skills)}):
{', '.join(missing_skills[:15]) if missing_skills else 'None'}

Based on this analysis, provide specific suggestions to improve the resume. Format your response as follows:

1. **Add Missing Skills** (if the candidate likely has relevant experience):
   - List 3-5 specific skills from the missing skills that should be added
   - For each skill, provide a brief bullet point example of how to describe it

2. **Strengthen Existing Content**:
   - Suggest 2-3 specific improvements to existing bullet points
   - Focus on quantifying achievements and using stronger action verbs

3. **Keyword Optimization**:
   - Suggest 2-3 industry-specific keywords to naturally incorporate

4. **Format & Structure**:
   - Any formatting improvements (if applicable)

Keep suggestions realistic - only suggest adding skills if they align with the candidate's existing experience. Be specific and actionable."""

        # Call Claude API
        message = client.messages.create(
            model=config.CLAUDE_MODEL,
            max_tokens=config.MAX_TOKENS,
            temperature=config.TEMPERATURE,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        return message.content[0].text
    
    except Exception as e:
        return f"Error getting AI suggestions: {str(e)}"

def generate_optimized_resume(resume_text, job_description, suggestions, user_approved_points):
    """Generate an optimized resume based on AI suggestions and user approval"""
    try:
        client = anthropic.Anthropic(api_key=config.CLAUDE_API_KEY)
        
        prompt = f"""You are an expert resume writer. Based on the original resume and approved suggestions, create an improved version.

ORIGINAL RESUME:
{resume_text}

JOB DESCRIPTION (for context):
{job_description}

AI SUGGESTIONS PROVIDED:
{suggestions}

USER APPROVED CHANGES:
{user_approved_points}

INSTRUCTIONS:
1. Keep the original structure and format as much as possible
2. Only implement the changes that the user approved
3. Maintain the candidate's authentic voice and experience
4. Do not fabricate experience or skills
5. Quantify achievements where possible
6. Use strong action verbs
7. Ensure ATS-friendly formatting

Return ONLY the improved resume text, maintaining the original format (no additional commentary)."""

        message = client.messages.create(
            model=config.CLAUDE_MODEL,
            max_tokens=config.MAX_TOKENS,
            temperature=0.5,  # Lower temperature for more consistent output
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        return message.content[0].text
    
    except Exception as e:
        return f"Error generating optimized resume: {str(e)}"

# ============================================================================
# RESUME DOCUMENT GENERATION
# ============================================================================

def create_resume_docx(resume_text, filename="optimized_resume.docx"):
    """Create a professional Word document from resume text"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(0.5)
        section.bottom_margin = docx.shared.Inches(0.5)
        section.left_margin = docx.shared.Inches(0.75)
        section.right_margin = docx.shared.Inches(0.75)
    
    # Split resume into paragraphs
    paragraphs = resume_text.split('\n')
    
    for para_text in paragraphs:
        if para_text.strip():
            paragraph = doc.add_paragraph(para_text)
            # Set font
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
    
    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

# ============================================================================
# FLASK ROUTES
# ============================================================================

@app.route('/')
def index():
    """Main page"""
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze():
    """Analyze resume-job match"""
    try:
        # Get job description
        job_description = request.form.get('job_description', '').strip()
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        # Get resume (either text or file)
        resume_text_input = request.form.get('resume_text', '').strip()
        resume_file = request.files.get('resume_file')
        
        # Process resume input
        resume_text = process_resume_input(resume_text_input, resume_file)
        
        if not resume_text:
            return jsonify({'error': 'Resume is required (either text or file)'}), 400
        
        # Run matching analysis
        print("\n🔍 Analyzing resume-job match...")
        results = inference_utils.match_resume_to_job_with_ml(resume_text, job_description)
        
        # We'll use the ML-based approach
        ml_results = results['approach_2_ml_based']
        
        # Prepare response
        response = {
            'success': True,
            'resume_text': resume_text,
            'job_description': job_description,
            'final_score': ml_results['final_score'],
            'bert_similarity': results['bert_similarity'],
            'tfidf_similarity': results['tfidf_similarity'],
            'category_match': results['category_match'],
            'resume_category': results['resume_category'],
            'job_category': results['job_category'],
            'skill_score': ml_results['skill_component_score'],
            'matched_skills': ml_results['matched_skills'],
            'missing_skills': ml_results['missing_skills'],
            'match_rate': ml_results['match_rate'],
            'recommendation': ml_results['recommendation'],
            'needs_optimization': ml_results['final_score'] < config.STRONG_MATCH_THRESHOLD
        }
        
        # Convert NumPy types to Python native types for JSON serialization
        response = convert_numpy_types(response)
        
        return jsonify(response)
    
    except Exception as e:
        print(f"❌ Error in analysis: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/get_suggestions', methods=['POST'])
def get_suggestions():
    """Get AI-powered improvement suggestions"""
    try:
        data = request.get_json()
        
        resume_text = data.get('resume_text')
        job_description = data.get('job_description')
        missing_skills = data.get('missing_skills', [])
        matched_skills = data.get('matched_skills', [])
        score_details = data.get('score_details', {})
        
        print("\n🤖 Getting AI suggestions...")
        suggestions = get_ai_suggestions(
            resume_text, 
            job_description, 
            missing_skills, 
            matched_skills,
            score_details
        )
        
        return jsonify({
            'success': True,
            'suggestions': suggestions
        })
    
    except Exception as e:
        print(f"❌ Error getting suggestions: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/optimize_resume', methods=['POST'])
def optimize_resume():
    """Generate optimized resume based on approved suggestions"""
    try:
        data = request.get_json()
        
        resume_text = data.get('resume_text')
        job_description = data.get('job_description')
        suggestions = data.get('suggestions')
        approved_points = data.get('approved_points', '')
        
        print("\n📝 Generating optimized resume...")
        optimized_text = generate_optimized_resume(
            resume_text,
            job_description,
            suggestions,
            approved_points
        )
        
        # Create DOCX file
        file_stream = create_resume_docx(optimized_text)
        
        return send_file(
            file_stream,
            as_attachment=True,
            download_name='optimized_resume.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        print(f"❌ Error optimizing resume: {str(e)}")
        return jsonify({'error': str(e)}), 500

# ============================================================================
# APP STARTUP
# ============================================================================

if __name__ == '__main__':
    # Load models before starting the app
    load_models()
    
    # Start Flask app
    print("\n🚀 Starting Resume Matcher Web Application...")
    print("📍 Access the app at: http://127.0.0.1:5001")
    print("\nPress CTRL+C to stop the server\n")
    
    app.run(debug=True, host='0.0.0.0', port=5001)